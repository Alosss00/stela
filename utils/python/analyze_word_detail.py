"""
Script untuk membaca detail lengkap dari satu file Word
"""
import os
from docx import Document
from docx.shared import Pt, Cm

def analyze_document_detailed(file_path):
    """Analisis detail dokumen Word"""
    print(f"\n{'='*100}")
    print(f"ANALISIS LENGKAP: {os.path.basename(file_path)}")
    print(f"{'='*100}\n")
    
    doc = Document(file_path)
    
    # HEADER DETAIL
    print("╔" + "═"*98 + "╗")
    print("║" + " "*40 + "HEADER SECTION" + " "*44 + "║")
    print("╚" + "═"*98 + "╝\n")
    
    for section in doc.sections:
        header = section.header
        
        # Cek tabel header
        for table in header.tables:
            print(f"📊 Header Table: {len(table.rows)} rows × {len(table.columns)} columns\n")
            for i, row in enumerate(table.rows):
                print(f"Row {i}: ")
                for j, cell in enumerate(row.cells):
                    if cell.text.strip():
                        print(f"  📍 Cell[{i},{j}]:")
                        for para in cell.paragraphs:
                            if para.text.strip():
                                print(f"     Text: '{para.text}'")
                                print(f"     Alignment: {para.alignment}")
                                for run in para.runs:
                                    font_size = run.font.size.pt if run.font.size else "Default"
                                    print(f"     Font: {run.font.name or 'Default'} | Size: {font_size}pt | Bold: {run.font.bold} | Italic: {run.font.italic}")
                print()
    
    # BODY CONTENT LENGKAP
    print("\n╔" + "═"*98 + "╗")
    print("║" + " "*40 + "BODY CONTENT" + " "*46 + "║")
    print("╚" + "═"*98 + "╝\n")
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            text_preview = para.text[:80] + "..." if len(para.text) > 80 else para.text
            print(f"\n📄 Paragraph {i+1}:")
            print(f"   Text: {text_preview}")
            print(f"   Alignment: {para.alignment}")
            print(f"   Style: {para.style.name}")
            
            # Detail runs
            if para.runs:
                run = para.runs[0]
                font_size = run.font.size.pt if run.font.size else "Default"
                print(f"   Font: {run.font.name or 'Default'} | Size: {font_size}pt | Bold: {run.font.bold}")
    
    # FOOTER DETAIL
    print("\n\n╔" + "═"*98 + "╗")
    print("║" + " "*40 + "FOOTER SECTION" + " "*44 + "║")
    print("╚" + "═"*98 + "╝\n")
    
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            if para.text.strip():
                print(f"📝 Footer Text: '{para.text}'")
                print(f"   Alignment: {para.alignment}")
                for run in para.runs:
                    font_size = run.font.size.pt if run.font.size else "Default"
                    print(f"   Font: {run.font.name or 'Default'} | Size: {font_size}pt | Italic: {run.font.italic}")

# Analisis file Pengawas Operasional sebagai contoh utama
file_path = r"assets\Draft Surat Penunjukan\Draft Contoh Surat Penunjukan Pengawas Operasional Pertambangan.docx"

if os.path.exists(file_path):
    analyze_document_detailed(file_path)
else:
    print(f"File tidak ditemukan: {file_path}")

print("\n\n" + "="*100)
print("✅ ANALISIS SELESAI")
print("="*100)
