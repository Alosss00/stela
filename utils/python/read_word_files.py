"""
Script untuk membaca file Word (.docx) dan mengekstrak informasi header, footer, dan konten
"""
import os
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def read_word_file(file_path):
    """Membaca file Word dan mengekstrak semua informasi"""
    print(f"\n{'='*80}")
    print(f"FILE: {os.path.basename(file_path)}")
    print(f"{'='*80}\n")
    
    doc = Document(file_path)
    
    # 1. HEADER
    print("=" * 80)
    print("HEADER SECTION")
    print("=" * 80)
    for section in doc.sections:
        header = section.header
        print(f"\n--- Header Content ---")
        for paragraph in header.paragraphs:
            if paragraph.text.strip():
                print(f"Text: {paragraph.text}")
                print(f"  Alignment: {paragraph.alignment}")
                print(f"  Style: {paragraph.style.name}")
                for run in paragraph.runs:
                    print(f"  Font: {run.font.name}, Size: {run.font.size}, Bold: {run.font.bold}, Italic: {run.font.italic}")
        
        # Cek tabel di header (biasanya untuk layout logo + text)
        for table in header.tables:
            print(f"\n--- Header Table Found ({len(table.rows)} rows x {len(table.columns)} cols) ---")
            for i, row in enumerate(table.rows):
                print(f"Row {i}:")
                for j, cell in enumerate(row.cells):
                    if cell.text.strip():
                        print(f"  Cell [{i},{j}]: {cell.text.strip()}")
    
    # 2. FOOTER
    print("\n" + "=" * 80)
    print("FOOTER SECTION")
    print("=" * 80)
    for section in doc.sections:
        footer = section.footer
        print(f"\n--- Footer Content ---")
        for paragraph in footer.paragraphs:
            if paragraph.text.strip():
                print(f"Text: {paragraph.text}")
                print(f"  Alignment: {paragraph.alignment}")
                print(f"  Style: {paragraph.style.name}")
                for run in paragraph.runs:
                    print(f"  Font: {run.font.name}, Size: {run.font.size}, Bold: {run.font.bold}, Italic: {run.font.italic}")
    
    # 3. BODY CONTENT (15 paragraf pertama)
    print("\n" + "=" * 80)
    print("BODY CONTENT (First 15 paragraphs)")
    print("=" * 80)
    for i, para in enumerate(doc.paragraphs[:15]):
        if para.text.strip():
            print(f"\nPara {i+1}: {para.text[:100]}...")
            print(f"  Alignment: {para.alignment}")
            print(f"  Style: {para.style.name}")
    
    # 4. PAGE SETUP
    print("\n" + "=" * 80)
    print("PAGE SETUP")
    print("=" * 80)
    section = doc.sections[0]
    print(f"Page Width: {section.page_width.cm:.2f} cm")
    print(f"Page Height: {section.page_height.cm:.2f} cm")
    print(f"Margin Top: {section.top_margin.cm:.2f} cm")
    print(f"Margin Bottom: {section.bottom_margin.cm:.2f} cm")
    print(f"Margin Left: {section.left_margin.cm:.2f} cm")
    print(f"Margin Right: {section.right_margin.cm:.2f} cm")
    
    print("\n" + "=" * 80 + "\n")

# Baca semua file Word di folder Draft Surat Penunjukan
draft_folder = r"assets\Draft Surat Penunjukan"

if os.path.exists(draft_folder):
    word_files = [f for f in os.listdir(draft_folder) if f.endswith('.docx') and not f.startswith('~')]
    
    print(f"\n\nMembaca {len(word_files)} file Word...\n")
    
    # Baca hanya 2 file pertama untuk contoh
    for i, filename in enumerate(word_files[:3]):
        file_path = os.path.join(draft_folder, filename)
        try:
            read_word_file(file_path)
        except Exception as e:
            print(f"Error reading {filename}: {str(e)}")
        
        if i < 2:  # Pause antara file
            print("\n" + "█" * 80)
            print("█" * 80 + "\n")
else:
    print(f"Folder tidak ditemukan: {draft_folder}")

print("\n✅ Selesai membaca file Word!")
